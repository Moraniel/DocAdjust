from pdf2docx import Converter
from pathlib import Path
from flask import Flask, request, render_template, send_file
from auth.drive_dowload import dowload_file
from auxiliary.leituraEscritaWord import read_file
from auxiliary.rename import rename
from docx.shared import Cm
from io import BytesIO
import pandas as pd
import docx 
import tempfile

app = Flask(__name__)
processed_file = None

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    # Cria um objeto Document
    document = docx.Document()

    # Define as margens do documento
    margins = {"top_margin": Cm(1), "bottom_margin": Cm(1), "left_margin": Cm(1), "right_margin": Cm(1)}
    for section in document.sections:
        for margin, value in margins.items():
            setattr(section, margin, value)

    if request.method == 'POST':
        # Lê o arquivo enviado pelo formulário
        file = request.files.get('file')
        if not file:
            return "Erro: nenhum arquivo enviado."

        # Lê o arquivo com base no seu tipo
        try:
            if file.content_type == 'text/csv':
                dados = pd.read_csv(BytesIO(file.read()))
            elif file.content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                dados = pd.read_excel(BytesIO(file.read()))
            elif file.content_type == 'application/octet-stream':
                dados = pd.read_excel(BytesIO(file.read()), engine="odf")
            else:
                raise Exception(f'Tipo de arquivo "{file.content_type}" não suportado.')
        except Exception as e:
            print(e)
            return "Erro ao processar o arquivo."

        # Converte os dados em uma lista de dicionários
        dados_dict = dados.to_dict(orient="records")
        
        # Renomeia as chaves dos dicionários
        rename(dados_dict)

        # Processa cada dicionário
        for i in dados_dict:
            # Remove a chave "Carimbo de data/hora"
            i.pop("Carimbo de data/hora", None)

            # Adiciona um título para cada resposta do setor
            document.add_heading(f"Respostas da {i['Identificação da Unidade/Gerência:']}", 0)

            # Processa cada chave e valor do dicionário
            for k, v in i.items():
                # Verifica se o valor não é NaN
                if not pd.isna(v):
                    document.add_heading(f"{k}", 1)
                    document.add_paragraph(f"{v}")

                    # Verifica se a chave é "Arquivo em PDF ou Documento (word ou odf)"
                    if k == "Arquivo em PDF ou Documento (word ou odf)":
                        try:
                            document.add_heading(f"Conteúdo do arquivo anexado pela {i['Identificação da Unidade/Gerência:']}", 1)
                            file_content = dowload_file(str(v))

                            # Verifica a extensão do arquivo
                            file_extension = Path(file_content).suffix.lower()

                            if file_extension == ".pdf":
                                # Converte o PDF em um arquivo Word
                                cv = Converter(file_content)
                                cv.convert("doc.docx")      
                                cv.close()
                                parent = docx.Document("doc.docx")
                            elif file_extension == '.docx':
                                # Lê o arquivo Word
                                parent = docx.Document(file_content)

                            # Adiciona o conteúdo do arquivo ao documento principal
                            read_file(parent, document)
                        except Exception as e:
                            return "Erro ao realizar a autenticação"
        
        # Salva o documento em um arquivo temporário e envia para o usuário
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as output:
            document.save(output)
            return send_file(output.name)
        
    return render_template('index.html')


@app.route('/download')
def download_file():
    global processed_file
    
    if processed_file is not None:
        try:
            return send_file(BytesIO(processed_file), as_attachment=True)
        except Exception as e:
            print(e)
            return "Erro ao fazer download do arquivo."
    else:
        return "Nenhum arquivo processado encontrado."
    
if __name__ == '__main__':
    app.run(debug=True)